import pdfplumber
from docx import Document
import logging
import io

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path_or_buffer):
    """Extract text from PDF using pdfplumber"""
    try:
        text = ""
        with pdfplumber.open(file_path_or_buffer) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {str(e)}")
                    continue
        
        result = text.strip()
        logger.info(f"Extracted {len(result)} characters from PDF")
        return result
        
    except Exception as e:
        logger.error(f"Error extracting PDF: {str(e)}")
        return ""

def extract_text_from_docx(file_path_or_buffer):
    """Extract text from DOCX using python-docx"""
    try:
        # Handle both file paths and BytesIO objects
        if isinstance(file_path_or_buffer, io.BytesIO):
            doc = Document(file_path_or_buffer)
        else:
            doc = Document(file_path_or_buffer)
            
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        result = text.strip()
        logger.info(f"Extracted {len(result)} characters from DOCX")
        return result
        
    except Exception as e:
        logger.error(f"Error extracting DOCX: {str(e)}")
        return ""

def extract_text_from_txt(file_content: bytes, encoding: str = 'utf-8'):
    """Extract text from TXT file with encoding detection"""
    try:
        # Try UTF-8 first
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Try other common encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for enc in encodings:
                try:
                    text = file_content.decode(enc)
                    logger.info(f"Successfully decoded with {enc}")
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # If all fail, use utf-8 with error replacement
                text = file_content.decode('utf-8', errors='replace')
                logger.warning("Used UTF-8 with error replacement")
        
        result = text.strip()
        logger.info(f"Extracted {len(result)} characters from TXT")
        return result
        
    except Exception as e:
        logger.error(f"Error extracting TXT: {str(e)}")
        return ""